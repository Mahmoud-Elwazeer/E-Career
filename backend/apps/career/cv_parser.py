"""
CV Parsing Service

This module provides CV parsing functionality using pdfplumber, python-docx, and docling.
It extracts structured data from CV files and maps skills to ESCO taxonomy.
"""

import json
import logging
import re
from pathlib import Path
from typing import Dict, List, Optional, Any
from django.conf import settings

from apps.skills.models import Skill
from apps.career.models import CareerUserSkill
from django.contrib.auth import get_user_model

User = get_user_model()

logger = logging.getLogger(__name__)


class CVParserService:
    """
    Service for parsing CV files and extracting structured data.
    
    Supports PDF, DOCX, and image formats (with OCR).
    """
    
    def __init__(self):
        self._bedrock = None
        self._model_id = getattr(settings, 'BEDROCK_MODEL_ID', 'anthropic.claude-sonnet-4-20250514-v1:0')
    
    @property
    def bedrock(self):
        """Lazy initialization of Bedrock client."""
        if self._bedrock is None:
            try:
                from apps.intelligence.career_ai import CareerAIService as BedrockService
                self._bedrock = BedrockService()
            except Exception as e:
                logger.warning(f"Failed to initialize Bedrock client: {e}")
                self._bedrock = None
        return self._bedrock
    
    def parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Parse PDF file and extract text.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            dict: Extracted text and metadata
        """
        try:
            import pdfplumber
            
            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"PDF file not found: {file_path}")
            
            text_parts = []
            metadata = {
                'file_name': path.name,
                'file_size': path.stat().st_size,
                'pages': 0,
            }
            
            with pdfplumber.open(path) as pdf:
                metadata['pages'] = len(pdf.pages)
                
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"--- Page {i + 1} ---\n{page_text}")
            
            return {
                'text': '\n\n'.join(text_parts),
                'metadata': metadata,
                'source': 'pdfplumber',
            }
            
        except ImportError:
            logger.error("pdfplumber not installed. Run: pip install pdfplumber")
            raise
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            raise
    
    def parse_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Parse DOCX file and extract text.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            dict: Extracted text and metadata
        """
        try:
            from docx import Document
            
            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"DOCX file not found: {file_path}")
            
            doc = Document(path)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            return {
                'text': '\n\n'.join(text_parts),
                'metadata': {
                    'file_name': path.name,
                    'file_size': path.stat().st_size,
                    'paragraphs': len(doc.paragraphs),
                    'tables': len(doc.tables),
                },
                'source': 'python-docx',
            }
            
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            raise
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            raise
    
    def parse_image(self, file_path: str) -> Dict[str, Any]:
        """
        Parse image file with OCR using docling.
        
        Args:
            file_path: Path to image file
            
        Returns:
            dict: Extracted text and metadata
        """
        try:
            from docling.document_converter import DocumentConverter
            
            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"Image file not found: {file_path}")
            
            converter = DocumentConverter()
            result = converter.convert(path)
            
            return {
                'text': result.document.export_to_text(),
                'metadata': {
                    'file_name': path.name,
                    'file_size': path.stat().st_size,
                    'format': path.suffix.lower(),
                },
                'source': 'docling',
            }
            
        except ImportError:
            logger.error("docling not installed. Run: pip install docling")
            raise
        except Exception as e:
            logger.error(f"Error parsing image {file_path}: {e}")
            raise
    
    def extract_structured_data(self, raw_text: str) -> Dict[str, Any]:
        """
        Extract structured data from raw CV text using AWS Bedrock Claude.
        
        Args:
            raw_text: Raw extracted text from CV
            
        Returns:
            dict: Structured CV data
        """
        if not self.bedrock or not self.bedrock.is_available:
            logger.warning("Bedrock not available, using fallback extraction")
            return self._fallback_extract_structured_data(raw_text)
        
        system_prompt = """You are an expert CV parser. Extract structured information from the provided CV text.

Return a JSON object with the following structure:
{
  "name": "Full name",
  "email": "email@example.com",
  "phone": "phone number",
  "location": "City, Country",
  "summary": "Professional summary (2-3 sentences)",
  "experience": [
    {
      "title": "Job title",
      "company": "Company name",
      "start_date": "YYYY-MM or YYYY",
      "end_date": "YYYY-MM or YYYY or 'Present'",
      "current": false,
      "description": "Job description and achievements"
    }
  ],
  "education": [
    {
      "degree": "Degree type and field",
      "institution": "Institution name",
      "start_date": "YYYY-MM or YYYY",
      "end_date": "YYYY-MM or YYYY or 'Expected'",
      "current": false
    }
  ],
  "skills": ["skill1", "skill2", "skill3"],
  "languages": ["language1", "language2"],
  "certifications": ["certification1", "certification2"]
}

Important:
- Extract all available information
- Use null for missing fields
- Normalize date formats to YYYY-MM
- Set current=true for current positions
- Extract skills as a flat list of skill names
- Only include skills that are clearly mentioned
"""
        
        prompt = f"""Parse this CV and extract structured information:

{raw_text[:15000]}  # Limit to avoid token overflow

Return ONLY the JSON object, no additional text."""
        
        try:
            response = self.bedrock.invoke_model(
                prompt=prompt,
                system_prompt=system_prompt,
                max_tokens=8000,
                temperature=0.1,
            )
            
            # Extract JSON from response
            json_start = response.find('{')
            json_end = response.rfind('}') + 1
            json_str = response[json_start:json_end]
            
            parsed_data = json.loads(json_str)
            return parsed_data
            
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse JSON from Bedrock response: {e}")
            logger.error(f"Response: {response[:500] if 'response' in dir() else 'N/A'}")
            return self._fallback_extract_structured_data(raw_text)
        except Exception as e:
            logger.error(f"Error extracting structured data: {e}")
            return self._fallback_extract_structured_data(raw_text)
    
    def _fallback_extract_structured_data(self, raw_text: str) -> Dict[str, Any]:
        """
        Fallback extraction using regex patterns when AI is unavailable.
        """
        data = {
            'name': None,
            'email': None,
            'phone': None,
            'location': None,
            'summary': None,
            'experience': [],
            'education': [],
            'skills': [],
            'languages': [],
            'certifications': [],
        }
        
        # Email
        email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', raw_text)
        if email_match:
            data['email'] = email_match.group(0)
        
        # Phone
        phone_match = re.search(r'[\+]?[(]?[0-9]{3}[)]?[-\s\.]?[0-9]{3}[-\s\.]?[0-9]{4,6}', raw_text)
        if phone_match:
            data['phone'] = phone_match.group(0)
        
        # Location (common patterns)
        location_patterns = [
            r'Location:\s*([^\n]+)',
            r'Address:\s*([^\n]+)',
            r'Based in\s*([^\n,]+)',
        ]
        for pattern in location_patterns:
            match = re.search(pattern, raw_text, re.IGNORECASE)
            if match:
                data['location'] = match.group(1).strip()
                break
        
        # Skills (common technical skills)
        common_skills = [
            'Python', 'JavaScript', 'Java', 'C++', 'C#', 'Ruby', 'PHP', 'Go', 'Rust',
            'HTML', 'CSS', 'SQL', 'React', 'Angular', 'Vue', 'Node.js', 'Django', 'Flask',
            'Docker', 'Kubernetes', 'AWS', 'Azure', 'GCP', 'Git', 'Linux', 'Redis', 'MongoDB',
            'TensorFlow', 'PyTorch', 'Scikit-learn', 'Pandas', 'NumPy', 'Excel', 'Power BI',
            'Machine Learning', 'Deep Learning', 'Data Analysis', 'Project Management',
            'Agile', 'Scrum', 'Communication', 'Leadership', 'Problem Solving', 'Teamwork',
        ]
        found_skills = []
        for skill in common_skills:
            if re.search(r'\b' + skill + r'\b', raw_text, re.IGNORECASE):
                found_skills.append(skill)
        data['skills'] = list(set(found_skills))[:20]  # Limit to 20 skills
        
        return data
    
    def map_skills_to_esco(self, skills: List[str]) -> List[Dict[str, Any]]:
        """
        Match extracted skill names to ESCO skills in the database using fuzzy matching.
        
        Args:
            skills: List of skill names to match
            
        Returns:
            list: List of matched skills with metadata
        """
        from difflib import SequenceMatcher
        
        matched_skills = []
        
        for skill_name in skills:
            if not skill_name:
                continue
            
            # Get all ESCO skills from database
            esco_skills = Skill.objects.filter(
                esco_uri__isnull=False
            ).values('id', 'esco_uri', 'name', 'name_ar', 'type', 'category')[:1000]
            
            best_match = None
            best_confidence = 0.0
            
            for esco_skill in esco_skills:
                # Calculate similarity between skill name and ESCO skill name
                confidence = SequenceMatcher(
                    None,
                    skill_name.lower(),
                    esco_skill['name'].lower()
                ).ratio()
                
                if confidence > best_confidence:
                    best_confidence = confidence
                    best_match = esco_skill
            
            # Only include matches with confidence > 0.7
            if best_match and best_confidence >= 0.7:
                matched_skills.append({
                    'skill_name': skill_name,
                    'esco_skill_id': best_match['id'],
                    'esco_uri': best_match['esco_uri'],
                    'esco_name': best_match['name'],
                    'confidence': best_confidence,
                    'esco_type': best_match['type'],
                    'esco_category': best_match['category'],
                })
        
        return matched_skills
    
    def update_user_skills(self, user: User, matched_skills: List[Dict[str, Any]]) -> int:
        """
        Update user's CareerUserSkill records from matched skills.
        
        Args:
            user: User instance
            matched_skills: List of matched skills from map_skills_to_esco
            
        Returns:
            int: Number of skills updated/created
        """
        updated_count = 0
        
        for match in matched_skills:
            try:
                skill = Skill.objects.get(id=match['esco_skill_id'])
                
                # Update or create CareerUserSkill
                user_skill, created = CareerUserSkill.objects.update_or_create(
                    user=user,
                    skill=skill,
                    defaults={
                        'proficiency': 'intermediate',
                        'years_experience': 0,
                        'verified': False,
                        'verification_source': 'cv_extraction',
                        'source': 'cv_extraction',
                        'confidence': match['confidence'],
                    }
                )
                
                if created or user_skill.confidence != match['confidence']:
                    updated_count += 1
                    
            except Skill.DoesNotExist:
                logger.warning(f"Skill not found: {match['esco_skill_id']}")
        
        return updated_count


# Singleton instance
cv_parser_service = CVParserService()