"""
Unified Document Processing Service.

Uses Docling (MIT, IBM) as the primary extraction engine,
with fallback to pdfplumber for simple PDFs.
Handles: PDF, DOCX, images (OCR), and text extraction
with structured output suitable for AI processing.
"""
from __future__ import annotations

import structlog
from pathlib import Path
from typing import Any

logger = structlog.get_logger()


class DocumentProcessor:
    """Unified document processing using Docling."""

    def __init__(self):
        self._converter = None

    @property
    def converter(self):
        if self._converter is None:
            try:
                from docling.document_converter import DocumentConverter
                self._converter = DocumentConverter()
            except ImportError:
                logger.warning("docling_not_installed", fallback="pdfplumber")
                self._converter = "fallback"
        return self._converter

    def extract_text(self, file_path: str | Path) -> DocumentResult:
        """Extract text and structure from a document."""
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()

        if self.converter != "fallback":
            return self._extract_with_docling(file_path)

        if suffix == ".pdf":
            return self._extract_pdf_fallback(file_path)
        elif suffix in (".docx", ".doc"):
            return self._extract_docx_fallback(file_path)
        elif suffix in (".txt", ".md"):
            return self._extract_text_file(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    def extract_from_bytes(self, content: bytes, filename: str) -> DocumentResult:
        """Extract from file content bytes."""
        import tempfile
        suffix = Path(filename).suffix
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as f:
            f.write(content)
            f.flush()
            return self.extract_text(f.name)

    def _extract_with_docling(self, file_path: Path) -> DocumentResult:
        """Primary extraction using Docling."""
        try:
            from docling.document_converter import DocumentConverter

            result = self.converter.convert(str(file_path))
            doc = result.document

            markdown_text = doc.export_to_markdown()

            tables = []
            if hasattr(doc, 'tables'):
                for table in doc.tables:
                    tables.append(table.export_to_dataframe().to_dict() if hasattr(table, 'export_to_dataframe') else {})

            metadata = {}
            if hasattr(doc, 'metadata'):
                metadata = doc.metadata if isinstance(doc.metadata, dict) else {}

            return DocumentResult(
                text=markdown_text,
                tables=tables,
                metadata=metadata,
                pages=getattr(doc, 'num_pages', 1),
                method="docling",
            )

        except Exception as e:
            logger.error("docling_extraction_failed", file=str(file_path), error=str(e))
            if file_path.suffix.lower() == ".pdf":
                return self._extract_pdf_fallback(file_path)
            raise

    def _extract_pdf_fallback(self, file_path: Path) -> DocumentResult:
        """Fallback PDF extraction using pdfplumber."""
        import pdfplumber

        text_parts = []
        tables = []
        num_pages = 0

        with pdfplumber.open(str(file_path)) as pdf:
            num_pages = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

                page_tables = page.extract_tables()
                if page_tables:
                    for table in page_tables:
                        tables.append(table)

        return DocumentResult(
            text="\n\n".join(text_parts),
            tables=tables,
            metadata={"file": str(file_path)},
            pages=num_pages,
            method="pdfplumber",
        )

    def _extract_docx_fallback(self, file_path: Path) -> DocumentResult:
        """Fallback DOCX extraction using python-docx."""
        from docx import Document

        doc = Document(str(file_path))
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                table_data.append([cell.text for cell in row.cells])
            tables.append(table_data)

        return DocumentResult(
            text="\n".join(text_parts),
            tables=tables,
            metadata={"file": str(file_path)},
            pages=1,
            method="python-docx",
        )

    def _extract_text_file(self, file_path: Path) -> DocumentResult:
        """Simple text file reading."""
        text = file_path.read_text(encoding="utf-8", errors="ignore")
        return DocumentResult(
            text=text,
            tables=[],
            metadata={"file": str(file_path)},
            pages=1,
            method="text",
        )


class DocumentResult:
    """Structured result from document extraction."""

    def __init__(
        self,
        text: str,
        tables: list[Any] = None,
        metadata: dict[str, Any] = None,
        pages: int = 1,
        method: str = "unknown",
    ):
        self.text = text
        self.tables = tables or []
        self.metadata = metadata or {}
        self.pages = pages
        self.method = method

    @property
    def word_count(self) -> int:
        return len(self.text.split())

    @property
    def is_empty(self) -> bool:
        return len(self.text.strip()) == 0

    def chunk(self, max_chars: int = 4000, overlap: int = 200) -> list[str]:
        """Split text into chunks suitable for LLM processing."""
        if len(self.text) <= max_chars:
            return [self.text]

        chunks = []
        start = 0
        while start < len(self.text):
            end = start + max_chars
            if end < len(self.text):
                split_at = self.text.rfind("\n", start, end)
                if split_at == -1 or split_at <= start:
                    split_at = self.text.rfind(" ", start, end)
                if split_at > start:
                    end = split_at

            chunks.append(self.text[start:end].strip())
            start = end - overlap

        return [c for c in chunks if c]

    def to_dict(self) -> dict[str, Any]:
        return {
            "text": self.text,
            "tables": self.tables,
            "metadata": self.metadata,
            "pages": self.pages,
            "method": self.method,
            "word_count": self.word_count,
        }


_processor: DocumentProcessor | None = None


def get_document_processor() -> DocumentProcessor:
    """Get singleton document processor."""
    global _processor
    if _processor is None:
        _processor = DocumentProcessor()
    return _processor
