"""
Resume Export Service - generates PDF/HTML/JSON from resume data.

Uses Django templates for HTML rendering and xhtml2pdf for PDF conversion.
"""
import io
import json
import logging
import os
from typing import Optional

from django.template.loader import render_to_string
from django.conf import settings

logger = logging.getLogger(__name__)


class ResumeExportService:
    TEMPLATES = {
        'modern': 'resume/export/modern.html',
        'professional': 'resume/export/professional.html',
        'creative': 'resume/export/creative.html',
        'minimalist': 'resume/export/minimalist.html',
    }

    def export_pdf(self, resume) -> Optional[bytes]:
        html = self._render_html(resume)
        try:
            from xhtml2pdf import pisa
            buffer = io.BytesIO()
            pisa_status = pisa.CreatePDF(io.StringIO(html), dest=buffer)
            if pisa_status.err:
                logger.error("PDF generation failed: %d errors", pisa_status.err)
                return None
            return buffer.getvalue()
        except ImportError:
            logger.warning("xhtml2pdf not installed, falling back to HTML export")
            return html.encode('utf-8')

    def export_html(self, resume) -> str:
        return self._render_html(resume)

    def export_docx(self, resume) -> Optional[bytes]:
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.error("python-docx not installed")
            return None

        doc = Document()
        style = doc.styles['Normal']
        style.font.size = Pt(10)
        style.font.name = 'Calibri'

        info = resume.personal_info or {}
        if info.get('full_name'):
            heading = doc.add_heading(info['full_name'], level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        contact_parts = [info.get(k) for k in ('email', 'phone', 'location', 'linkedin') if info.get(k)]
        if contact_parts:
            p = doc.add_paragraph(' | '.join(contact_parts))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.style.font.size = Pt(9)

        if resume.summary:
            doc.add_paragraph(resume.summary)

        if resume.experience:
            doc.add_heading('Experience', level=1)
            for exp in resume.experience:
                p = doc.add_paragraph()
                run = p.add_run(f"{exp.get('title', '')} at {exp.get('company', '')}")
                run.bold = True
                date_str = f"{exp.get('start_date', '')} — {'Present' if exp.get('current') else exp.get('end_date', '')}"
                p.add_run(f"  ({date_str})")
                if exp.get('description'):
                    doc.add_paragraph(exp['description'], style='List Bullet')

        if resume.education:
            doc.add_heading('Education', level=1)
            for edu in resume.education:
                degree = edu.get('degree', '')
                field = f" in {edu.get('field')}" if edu.get('field') else ''
                doc.add_paragraph(f"{degree}{field} — {edu.get('school', '')}")

        if resume.skills:
            doc.add_heading('Skills', level=1)
            doc.add_paragraph(', '.join(resume.skills))

        if resume.certifications:
            doc.add_heading('Certifications', level=1)
            for cert in resume.certifications:
                doc.add_paragraph(cert, style='List Bullet')

        if resume.languages:
            doc.add_heading('Languages', level=1)
            doc.add_paragraph(', '.join(resume.languages))

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def export_json(self, resume) -> str:
        return json.dumps({
            'personal_info': resume.personal_info,
            'summary': resume.summary,
            'experience': resume.experience,
            'education': resume.education,
            'skills': resume.skills,
            'projects': resume.projects,
            'certifications': resume.certifications,
            'languages': resume.languages,
            'interests': resume.interests,
        }, indent=2)

    def _render_html(self, resume) -> str:
        template_category = 'modern'
        if resume.template:
            template_category = resume.template.category

        template_name = self.TEMPLATES.get(template_category, self.TEMPLATES['modern'])

        try:
            return render_to_string(template_name, {'resume': resume})
        except Exception:
            return render_to_string('resume/export/modern.html', {'resume': resume})


resume_export_service = ResumeExportService()
