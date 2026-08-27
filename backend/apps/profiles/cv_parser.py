"""
CV file processing and text extraction with multiple parser plugins.

This module implements a plugin-based CV parsing system that:
1. Detects file type
2. Extracts text using the appropriate parser
3. Structures data via Haiku LLM
4. Extracts skills and maps to ESCO taxonomy
"""
import os
import logging
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple
from dataclasses import dataclass, field
from abc import ABC, abstractmethod

from django.core.files.uploadedfile import UploadedFile

logger = logging.getLogger(__name__)


# ============================================================================
# Parser Plugin Interface
# ============================================================================

class ResumeParserPlugin(ABC):
    """Abstract base class for CV parser plugins."""
    
    @abstractmethod
    def can_handle(self, file_path: str, file_type: str) -> bool:
        """Check if this parser can handle the given file."""
        pass
    
    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        """Extract text from the file."""
        pass
    
    @property
    @abstractmethod
    def name(self) -> str:
        """Name of this parser."""
        pass


# ============================================================================
# PDF Parsers
# ============================================================================

class DoclingParserPlugin(ResumeParserPlugin):
    """
    IBM Docling parser for PDF files.
    
    Uses Docling (MIT license) for high-quality PDF parsing with layout detection.
    """
    
    def __init__(self):
        self._client = None
    
    @property
    def name(self) -> str:
        return "docling"
    
    def can_handle(self, file_path: str, file_type: str) -> bool:
        return file_type == "pdf"
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from PDF using Docling."""
        try:
            from docling.document_converter import DocumentConverter
            
            converter = DocumentConverter()
            result = converter.convert(file_path)
            
            # Extract text from the document
            text = result.document.export_to_text()
            return text
            
        except ImportError:
            logger.warning("Docling not installed, falling back to pdfplumber")
            return self._fallback_to_pdfplumber(file_path)
        except Exception as e:
            logger.error(f"Docling extraction failed: {e}")
            return self._fallback_to_pdfplumber(file_path)
    
    def _fallback_to_pdfplumber(self, file_path: str) -> str:
        """Fallback to pdfplumber if Docling fails."""
        try:
            import pdfplumber
            
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"pdfplumber fallback failed: {e}")
            return ""


class PdfplumberParserPlugin(ResumeParserPlugin):
    """pdfplumber parser for PDF files - pure text extraction."""
    
    @property
    def name(self) -> str:
        return "pdfplumber"
    
    def can_handle(self, file_path: str, file_type: str) -> bool:
        return file_type == "pdf"
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from PDF using pdfplumber."""
        try:
            import pdfplumber
            
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"pdfplumber extraction failed: {e}")
            return ""


# ============================================================================
# Image/Scanned CV Parser
# ============================================================================

class EasyOCRParserPlugin(ResumeParserPlugin):
    """
    EasyOCR parser for scanned/image CVs.
    
    Uses EasyOCR for OCR on image-based CVs.
    """
    
    def __init__(self):
        self._reader = None
    
    @property
    def name(self) -> str:
        return "easyocr"
    
    def can_handle(self, file_path: str, file_type: str) -> bool:
        # Handle image files and PDFs (will convert to images)
        return file_type in ("png", "jpg", "jpeg", "gif", "bmp", "tiff", "pdf")
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from image/PDF using EasyOCR."""
        try:
            import easyocr
            from pdf2image import convert_from_path
            
            # Initialize reader (only once)
            if self._reader is None:
                self._reader = easyocr.Reader(['en'])
            
            # If PDF, convert to images first
            if file_path.lower().endswith('.pdf'):
                images = convert_from_path(file_path)
                text_parts = []
                for img in images:
                    result = self._reader.readtext(img)
                    for (bbox, text, prob) in result:
                        text_parts.append(text)
                return "\n\n".join(text_parts)
            else:
                # Direct image OCR
                result = self._reader.readtext(file_path)
                text_parts = [text for (bbox, text, prob) in result]
                return "\n\n".join(text_parts)
                
        except ImportError:
            logger.warning("EasyOCR not installed")
            return ""
        except Exception as e:
            logger.error(f"EasyOCR extraction failed: {e}")
            return ""


# ============================================================================
# Document Parsers
# ============================================================================

class DocxParserPlugin(ResumeParserPlugin):
    """python-docx parser for DOCX files."""
    
    @property
    def name(self) -> str:
        return "docx"
    
    def can_handle(self, file_path: str, file_type: str) -> bool:
        return file_type in ("docx", "doc")
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from DOCX."""
        try:
            import docx
            
            doc = docx.Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return ""


class TxtParserPlugin(ResumeParserPlugin):
    """Plain text parser for TXT files."""
    
    @property
    def name(self) -> str:
        return "txt"
    
    def can_handle(self, file_path: str, file_type: str) -> bool:
        return file_type == "txt"
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.read()
            except Exception as e:
                logger.error(f"TXT extraction failed: {e}")
                return ""
        except Exception as e:
            logger.error(f"TXT extraction failed: {e}")
            return ""


# ============================================================================
# Main CV Parser
# ============================================================================

@dataclass
class ParsedCVData:
    """Structured data from parsed CV."""
    personal: Dict[str, Any] = field(default_factory=dict)
    experience: List[Dict[str, Any]] = field(default_factory=list)
    education: List[Dict[str, Any]] = field(default_factory=list)
    skills: Dict[str, Any] = field(default_factory=dict)
    certifications: List[Dict[str, Any]] = field(default_factory=list)
    projects: List[Dict[str, Any]] = field(default_factory=list)
    raw_text: str = ""
    parser_used: str = ""


class CVParser:
    """Extract text from CV files with multiple parser support."""
    
    SUPPORTED_FORMATS = ['.pdf', '.docx', '.doc', '.txt', '.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff']
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    def __init__(self):
        self._plugins: List[ResumeParserPlugin] = []
        self._initialize_plugins()
    
    def _initialize_plugins(self):
        """Initialize all available parser plugins."""
        self._plugins = [
            DoclingParserPlugin(),
            PdfplumberParserPlugin(),
            EasyOCRParserPlugin(),
            DocxParserPlugin(),
            TxtParserPlugin(),
        ]
    
    def _get_file_type(self, file_path: str) -> str:
        """Get file type from extension."""
        ext = Path(file_path).suffix.lower().lstrip('.')
        return ext if ext else 'unknown'
    
    def _find_parser(self, file_path: str, file_type: str) -> Optional[ResumeParserPlugin]:
        """Find the appropriate parser for the file."""
        for plugin in self._plugins:
            if plugin.can_handle(file_path, file_type):
                return plugin
        return None
    
    def extract_text(self, file: UploadedFile) -> str:
        """
        Extract text from uploaded CV file.
        
        Args:
            file: Django UploadedFile instance
            
        Returns:
            str: Extracted text content
            
        Raises:
            ValueError: If file format not supported or file too large
        """
        # Validate file size
        if file.size > self.MAX_FILE_SIZE:
            raise ValueError(f"File too large. Maximum size is {self.MAX_FILE_SIZE / 1024 / 1024}MB")
        
        # Get file extension
        file_ext = Path(file.name).suffix.lower()
        file_type = self._get_file_type(file.name)
        
        if file_ext not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file format. Supported: {', '.join(self.SUPPORTED_FORMATS)}")
        
        # Save to temporary file with randomized name (prevents path traversal/collision)
        import uuid
        safe_ext = file_ext  # already validated against SUPPORTED_FORMATS
        temp_path = f"/tmp/cv_upload_{uuid.uuid4().hex}{safe_ext}"
        try:
            with open(temp_path, 'wb') as f:
                for chunk in file.chunks():
                    f.write(chunk)
            
            # Find and use appropriate parser
            parser = self._find_parser(temp_path, file_type)
            
            if parser is None:
                raise ValueError("No suitable parser found for this file type")
            
            text = parser.extract_text(temp_path)
            
            logger.info(
                "cv_text_extracted",
                parser=parser.name,
                file=file.name,
                text_length=len(text),
            )
            
            return text
            
        finally:
            # Clean up temp file
            if os.path.exists(temp_path):
                os.remove(temp_path)
    
    def parse_cv(self, file: UploadedFile) -> ParsedCVData:
        """
        Parse CV file and extract structured data.
        
        Args:
            file: Django UploadedFile instance
            
        Returns:
            ParsedCVData: Structured CV data
        """
        # Extract text
        raw_text = self.extract_text(file)
        
        # Get file type
        file_ext = Path(file.name).suffix.lower()
        file_type = self._get_file_type(file.name)
        
        # Find parser used
        parser = self._find_parser("", file_type)
        parser_name = parser.name if parser else "unknown"
        
        return ParsedCVData(
            raw_text=raw_text,
            parser_used=parser_name,
        )
    
    @classmethod
    def get_file_info(cls, file: UploadedFile) -> dict:
        """
        Get information about uploaded file.
        
        Args:
            file: Django UploadedFile instance
            
        Returns:
            dict: File information
        """
        file_ext = Path(file.name).suffix.lower()
        
        return {
            'name': file.name,
            'size': file.size,
            'extension': file_ext,
            'is_supported': file_ext in cls.SUPPORTED_FORMATS,
            'size_mb': round(file.size / (1024 * 1024), 2)
        }


# Singleton instance
cv_parser = CVParser()